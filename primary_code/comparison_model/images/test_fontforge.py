# -*- coding: utf-8 -*-

import docx

document = docx.Document("demo.docx")
from docx.shared import Pt

style = document.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(10)

Tstyle = document.styles['Normal']
font = Tstyle.font  # << this line assigns font = doc.styles['Normal'].font

#paragraph.style = document.styles['Normal']
text_name="Victor jagt zwölf Boxkämpfer quer über den großen Sylter Deic"
font.name = "Nunito Sans"
font.size = Pt(48)
doc=document


TStyle, FStyle = doc.styles['Normal'], doc.styles['Heading 1']
for style in (TStyle, FStyle):
    style.font.name = "Nunito Sans"
TStyle.font.size = Pt(48)
FStyle.font.size = Pt(24)


Title = doc.add_paragraph()
Title.style = TStyle
TRun = Title.add_run(text_name)
TRun.bold = True

FCreated = Title.add_run("the name ")

FLog = doc.add_paragraph()
FLog.style = FStyle
FinanceTitle = FLog.add_run("Log Begins:")

document.save('demo.docx')